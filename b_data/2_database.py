import docx
import openpyxl
import re

# 打开DOCX文件
# docx_file = docx.Document(".\\orignal_data\\kcjj\\9-xinanzhuanye-zhongwenjianjie.docx")
docx_file = docx.Document(".\\orignal_data\\kcjj\\8-jikezhuanye-zhongwenjianjie.docx")

# 创建一个新的XLSX工作簿
workbook = openpyxl.Workbook()
worksheet = workbook.active

# 添加表头
worksheet.append(["课程名称", "学分", "总学时", "先修课程", "简介"])

# 初始化变量用于存储课程信息和简介
course_name = ""
credit = ""
total_hours = ""
prerequisite_courses = ""
course_description = ""
found_description = False  # 用于标记是否找到课程简介

# 遍历DOCX文件中的段落
for paragraph in docx_file.paragraphs:
    # 检查段落中是否包含"课程名称："字段
    if "课程名称：" in paragraph.text:
        # 提取字段后面的值
        course_name = paragraph.text.split("：", 1)[1].strip()
        print(course_name)
    elif "总学时：" in paragraph.text:
        match = re.search(r"学分：\s*([\d.]+)\s+总学时：\s*([\d.]+)", paragraph.text)
        if match:
            credit = match.group(1)
            total_hours = match.group(2)
            print(credit)
            print(total_hours)
        if match == "": print("match:", match)
    elif paragraph.text.startswith("先修课程："):
        prerequisite_courses = paragraph.text.split("：")[1].strip()
        if prerequisite_courses == "":
            prerequisite_courses = "无"
        print(prerequisite_courses)
    elif paragraph.text.startswith("课程简介："):
        # 找到课程简介行，标记为找到课程简介
        found_description = True
    elif found_description:
        # 在找到课程简介后提取下一段文字
        course_description = paragraph.text.strip()
        found_description = False

        # 如果找到了所有课程信息，将其写入Excel文件
        if course_name and credit and total_hours and prerequisite_courses:
            worksheet.append([course_name, credit, total_hours, prerequisite_courses, course_description])

            # 重置变量以提取下一个课程信息
            course_name = ""
            credit = ""
            total_hours = ""
            prerequisite_courses = ""
            course_description = ""
        if found_description == True:print("found_description",found_description)
# 保存XLSX文件
# workbook.save("final_data/信安专业课程详情.xlsx")
workbook.save("final_data/计科专业课程详情.xlsx")